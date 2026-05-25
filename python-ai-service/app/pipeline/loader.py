"""
Document Loader — unified interface for PDF, DOCX, TXT, MD.

Uses LangChain loaders with proper metadata extraction.
"""
from pathlib import Path
from typing import Optional
from dataclasses import dataclass, field
from app.core.logger import get_logger

logger = get_logger(__name__)


@dataclass
class LoadedDocument:
    """Result of loading a document."""
    content: str
    file_name: str
    file_type: str
    size_bytes: int
    metadata: dict = field(default_factory=dict)


class DocumentLoader:
    """Load documents from PDF, DOCX, TXT, MD files.

    Uses LangChain community loaders. Falls back gracefully
    for absent optional dependencies.
    """

    CHUNK_SIZE_FOR_LOADING = 1000000  # 1MB chunks for large files

    def __init__(self):
        self._loaders: dict[str, callable] = {}
        self._register_builtins()

    def _register_builtins(self):
        self._loaders[".txt"] = self._load_text
        self._loaders[".md"] = self._load_text
        self._loaders[".csv"] = self._load_text

    def _load_text(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    async def load(self, file_path: str) -> LoadedDocument:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        content: str

        if ext == ".pdf":
            content = self._load_pdf(file_path)
        elif ext == ".docx":
            content = self._load_docx(file_path)
        elif ext in self._loaders:
            content = self._loaders[ext](file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}. Supported: pdf, docx, txt, md")

        return LoadedDocument(
            content=content,
            file_name=path.name,
            file_type=ext.lstrip("."),
            size_bytes=path.stat().st_size,
            metadata={"source": str(path.absolute())},
        )

    def _load_pdf(self, file_path: str) -> str:
        try:
            from langchain_community.document_loaders import PyPDFLoader
            docs = PyPDFLoader(file_path).load()
            parts = []
            for i, doc in enumerate(docs):
                parts.append(f"[Page {i + 1}]\n{doc.page_content}")
            text = "\n\n".join(parts)
            logger.info(f"PDF loaded: {len(docs)} pages, {len(text)} chars")
            return text
        except ImportError:
            return self._load_pdf_fallback(file_path)

    def _load_pdf_fallback(self, file_path: str) -> str:
        """Fallback using pypdf directly."""
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            parts = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    parts.append(f"[Page {i + 1}]\n{text}")
            return "\n\n".join(parts)
        except ImportError:
            raise ImportError("Install pypdf or langchain-community for PDF support")

    def _load_docx(self, file_path: str) -> str:
        try:
            from langchain_community.document_loaders import Docx2txtLoader
            docs = Docx2txtLoader(file_path).load()
            parts = []
            for doc in docs:
                parts.append(doc.page_content)
            text = "\n\n".join(parts)
            logger.info(f"DOCX loaded: {len(text)} chars")
            return text
        except ImportError:
            return self._load_docx_fallback(file_path)

    def _load_docx_fallback(self, file_path: str) -> str:
        try:
            from docx import Document
            doc = Document(file_path)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells if cell.text)
                    if row_text.strip():
                        parts.append(row_text)
            return "\n\n".join(parts)
        except ImportError:
            raise ImportError("Install python-docx or langchain-community for DOCX support")

    @property
    def supported_types(self) -> list[str]:
        return ["pdf", "docx", "txt", "md", "csv"]


document_loader = DocumentLoader()
