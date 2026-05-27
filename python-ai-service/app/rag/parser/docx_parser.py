"""DOCX Parser — 基于 python-docx 的企业级 Word 文档解析.

特性:
  - 段落提取（含样式检测 → 标题识别）
  - 表格文本提取
  - 标题层级识别 (Heading 1-6)
  - 空行/空白段落清洗
"""

from io import BytesIO
from app.rag.parser.parser_factory import BaseParser, ParsedDocument, ParagraphInfo, HeadingInfo
from app.core.logger import get_logger

logger = get_logger(__name__)


class DocxParser(BaseParser):
    """DOCX 文档解析器 — python-docx 后端."""

    # Word 内置 Heading 样式名
    HEADING_STYLES = frozenset({
        "Heading 1", "Heading 2", "Heading 3",
        "Heading 4", "Heading 5", "Heading 6",
        "heading 1", "heading 2", "heading 3",
        "1. Heading 1", "2. Heading 2", "3. Heading 3",
        "标题 1", "标题 2", "标题 3", "标题 4", "标题 5", "标题 6",
    })

    def _heading_level_from_style(self, style_name: str) -> int:
        """从样式名提取标题级别."""
        sl = style_name.lower()
        for i in range(1, 7):
            if f"heading {i}" in sl or f"标题 {i}" in sl or f"{i}. heading" in sl:
                return i
        return 0

    @property
    def supported_types(self) -> list[str]:
        return [".docx"]

    def parse(self, file_path: str, file_name: str, content: bytes) -> ParsedDocument:
        from docx import Document as DocxDocument

        doc = DocxDocument(BytesIO(content))
        paragraphs: list[ParagraphInfo] = []
        headings: list[HeadingInfo] = []
        text_parts: list[str] = []
        errors: list[str] = []
        char_offset = 0

        # 辅助: 遍历 doc.element.body 以按文档顺序处理段落和表格
        from docx.oxml.ns import qn
        body = doc.element.body

        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                # 段落
                para = _find_paragraph_by_element(doc, child)
                if para is None:
                    continue

                text = para.text or ""
                stripped = text.strip()
                if not stripped:
                    char_offset += len(text) + 1
                    continue

                style_name = para.style.name if para.style else ""
                heading_level = self._heading_level_from_style(style_name)
                is_heading = heading_level > 0

                paragraphs.append(ParagraphInfo(
                    text=stripped,
                    page_number=0,  # DOCX 无原生页码
                    char_offset=char_offset,
                    char_count=len(stripped),
                    is_heading=is_heading,
                    heading_level=heading_level,
                ))

                if is_heading:
                    headings.append(HeadingInfo(
                        level=heading_level,
                        text=stripped,
                        page_number=0,
                        char_offset=char_offset,
                    ))

                text_parts.append(stripped)
                char_offset += len(text) + 1

            elif tag == "tbl":
                # 表格
                table_texts = _extract_table_text(child)
                for row_text in table_texts:
                    if row_text.strip():
                        paragraphs.append(ParagraphInfo(
                            text=row_text,
                            page_number=0,
                            char_offset=char_offset,
                            char_count=len(row_text),
                            is_heading=False,
                            heading_level=0,
                        ))
                        text_parts.append(row_text)
                        char_offset += len(row_text) + 1

        full_text = "\n\n".join(text_parts)
        title_path = [h.text for h in headings if h.level <= 2]

        logger.info(
            f"[DocxParser] 解析完成: '{file_name}' "
            f"chars={len(full_text)} paras={len(paragraphs)} "
            f"headings={len(headings)}"
        )

        return ParsedDocument(
            file_name=file_name,
            file_type="docx",
            total_chars=len(full_text),
            total_pages=0,  # DOCX 无原生页码
            full_text=full_text,
            paragraphs=paragraphs,
            headings=headings,
            title_path=title_path,
            metadata={
                "parser": "python-docx",
                "paragraph_count": len(paragraphs),
                "table_count": sum(1 for c in body if c.tag.split("}")[-1] == "tbl"),
                "heading_count": len(headings),
            },
            parse_errors=errors,
        )


def _find_paragraph_by_element(doc, element):
    """在 doc.paragraphs 中找到对应的 Paragraph 对象."""
    for p in doc.paragraphs:
        if p._element is element:
            return p
    return None


def _extract_table_text(table_element) -> list[str]:
    """从表格 XML 元素中提取文本，每行一个字符串."""
    from docx.oxml.ns import qn
    rows = table_element.findall(qn("w:tr"))
    row_texts = []
    for row in rows:
        cells = row.findall(qn("w:tc"))
        cell_texts = []
        for cell in cells:
            paras = cell.findall(qn("w:p"))
            cell_text = " ".join(
                "".join(t.text or "" for t in p.findall(qn("w:t")))
                for p in paras
            ).strip()
            if cell_text:
                cell_texts.append(cell_text)
        if cell_texts:
            row_texts.append(" | ".join(cell_texts))
    return row_texts
